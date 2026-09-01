"""
Astrova Audit Report — DOCX Export Script
Converts the markdown audit report to a professionally formatted DOCX file.

Usage:
    python docs/audit/export_docx.py

Output:
    docs/audit/ASTROVA_MASTER_PROJECT_AUDIT_REPORT.docx
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

# ── Paths ──────────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(SCRIPT_DIR)
REPORT_MD = os.path.join(SCRIPT_DIR, "ASTROVA_MASTER_PROJECT_AUDIT_REPORT.md")
REPORT_DOCX = os.path.join(SCRIPT_DIR, "ASTROVA_MASTER_PROJECT_AUDIT_REPORT.docx")

# ── Brand Colors ───────────────────────────────────────────────────────────
TERRACOTTA = RGBColor(0xC2, 0x70, 0x3E)    # Primary brand
DARK_CHARCOAL = RGBColor(0x1A, 0x1A, 0x1A)  # Headings
STONE = RGBColor(0x57, 0x53, 0x4E)          # Body text
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF4)    # Table alt rows
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RED = RGBColor(0xDC, 0x26, 0x26)
GREEN = RGBColor(0x16, 0xA3, 0x4A)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)


def parse_markdown(md_text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Heading
        if line.startswith("# "):
            blocks.append(("h1", line[2:].strip()))
        elif line.startswith("## "):
            blocks.append(("h2", line[3:].strip()))
        elif line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
        
        # Table
        elif line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Will be incremented at end of loop
            blocks.append(("table", table_lines))
        
        # Horizontal rule
        elif line.strip() == "---":
            blocks.append(("hr", None))
        
        # Bold line
        elif line.startswith("**") and line.endswith("**"):
            blocks.append(("bold", line.strip("*").strip()))
        
        # Bullet list
        elif line.startswith("- "):
            blocks.append(("bullet", line[2:].strip()))
        
        # Numbered list
        elif re.match(r"^\d+\. ", line):
            blocks.append(("numbered", re.sub(r"^\d+\. ", "", line).strip()))
        
        # Regular paragraph
        elif line.strip():
            blocks.append(("p", line.strip()))
        
        i += 1
    
    return blocks


def add_formatted_text(paragraph, text):
    """Add text with inline formatting (bold, code, etc.)."""
    # Split by bold markers
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for idx, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        if idx % 2 == 1:  # Bold part
            run.bold = True
            run.font.color.rgb = DARK_CHARCOAL
        else:
            run.font.color.rgb = STONE
        run.font.size = Pt(10)


def create_docx():
    """Create the DOCX report."""
    # Read markdown
    with open(REPORT_MD, "r", encoding="utf-8") as f:
        md_text = f.read()
    
    blocks = parse_markdown(md_text)
    
    # Create document
    doc = Document()
    
    # ── Page Setup ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ── Styles ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = STONE
    
    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = DARK_CHARCOAL
    
    # ── Title Page ──────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ASTROVA")
    run.font.size = Pt(36)
    run.font.color.rgb = TERRACOTTA
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Master Project Audit Report")
    run.font.size = Pt(20)
    run.font.color.rgb = DARK_CHARCOAL
    
    doc.add_paragraph()
    
    # Metadata
    meta_lines = [
        "Date: August 31, 2026",
        "Project: Astrova (formerly Dharohar AI / Heritage Atlas)",
        "Audit Type: Comprehensive Technical Audit",
        "Auditor: Buffy (Codebuff AI Agent)",
        "Model: MiMo 2.5 Balanced",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = STONE
    
    # Page break
    doc.add_page_break()
    
    # ── Process Blocks ──────────────────────────────────────────────────
    for block_type, content in blocks:
        
        if block_type == "h1":
            doc.add_page_break()
            h = doc.add_heading(content, level=1)
            for run in h.runs:
                run.font.color.rgb = TERRACOTTA
        
        elif block_type == "h2":
            doc.add_paragraph()
            h = doc.add_heading(content, level=2)
            for run in h.runs:
                run.font.color.rgb = DARK_CHARCOAL
        
        elif block_type == "h3":
            h = doc.add_heading(content, level=3)
            for run in h.runs:
                run.font.color.rgb = TERRACOTTA
        
        elif block_type == "table":
            table_lines = content
            if len(table_lines) < 2:
                continue
            
            # Parse header
            headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
            
            # Skip separator line
            data_lines = table_lines[2:] if len(table_lines) > 2 else []
            
            # Parse data rows
            rows = []
            for line in data_lines:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                if cells:
                    rows.append(cells)
            
            if not headers:
                continue
            
            # Create table
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Header row
            for j, header in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(header)
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE
                # Background
                from docx.oxml.ns import qn
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn("w:shd"), {
                    qn("w:fill"): "C2703E",
                    qn("w:val"): "clear"
                })
                shading.append(shading_elm)
            
            # Data rows
            for i, row in enumerate(rows):
                for j, cell_text in enumerate(row):
                    if j >= len(headers):
                        continue
                    cell = table.rows[i + 1].cells[j]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(9)
                    run.font.color.rgb = STONE
               
